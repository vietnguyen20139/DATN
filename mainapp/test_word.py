import json
from docx import Document
import pandas as pd
import connect_firebase
import os

# def create_table_from_json(json_data, output_file):
    # Create a new Word document

    # Load JSON data


def createDox(orderID):
    data = connect_firebase.loadDataFromFirebase(f"orders/{orderID}")

    doc = Document()

    doc.add_heading('Thông tin đơn hàng', 0) 


    # print(data.val())
    dataVal = data.val()

    keys = ["orderID", "dateCreateOrder", "createPerson", "phoneCreatePerson","cargo1", "getDay", "typeContainer", "carTon", "driver", "car", "distanceRunning", "timeRunning","maxHum", "minHum", "maxTemp", "minTemp"]
    keys_name = ["Mã đơn", "Ngày tạo", "Số điện thoại", "Người tạo","Hàng hóa", "Ngày lấy hàng", "Loại xe", "Trọng tải (tấn)", "Tài xế", "Biển số xe", "Quảng đường (km)", "Thời gian di chuyển (Tiếng)", "Độ ẩm tối đa (%)", "Độ ẩm tối thiểu (%)", "Nhiệt độ tối đa (℃)", "Nhiệt độ tối thiểu (℃)"]
   
    table = doc.add_table(rows=len(keys), cols=2)

    for row, (order_id, order_info) in enumerate(dataVal.items()):
        # print(order_id)
        # print(order_info)
        
        for key_index, key in enumerate(keys):
            if key == order_id:
                value = order_info
                name_vietsub = keys_name[key_index]
                print(f"{keys_name[key_index]}: {value}")
                
                cell1 = table.cell(key_index, 0)  # Access the cell in the first row, second column
                cell2 = table.cell(key_index, 1)  # Access the cell in the first row, second column

                cell1.text = name_vietsub
                cell2.text = str(value)
                
    table.style = 'TableGrid'



    keys_2 = ["address", "location", "name", "phone", "note", "typeCargo"]
    keys_name_2 = ["Địa chỉ", "Tọa độ", "Tên khách hàng", "Số điện thoại","Ghi chú", "Phương thức lấy/trả hàng"]

    # doc.add_heading('Điểm dừng 1', 0) 


    numstop = data.val()["numsStop"]
    print(numstop)

    if numstop == '':
        print("numstop is null")
    
    if numstop:
        for i in range(int(numstop)):
                
            doc.add_heading(f"Thông tin điểm dừng {i + 1}", 0) 
            print(i)

            table = doc.add_table(rows=len(keys_2), cols=2)

            for row, (order_id, order_info) in enumerate(dataVal.items()):
                # print(order_id)
                # print(order_info)
                
                for key_index, key in enumerate(keys_2):
                    if (key + str(i + 1)) == order_id:
                        value = order_info
                        
                        if key == "typeCargo":
                            value = "lấy hàng" if order_info == "get" else "trả hàng"
                        name_vietsub = keys_name_2[key_index]
                        print(f"{keys_name_2[key_index]}: {value}")
                        
                        cell1 = table.cell(key_index, 0)  # Access the cell in the first row, second column
                        cell2 = table.cell(key_index, 1)  # Access the cell in the first row, second column

                        cell1.text = name_vietsub
                        cell2.text = str(value)
                        
            table.style = 'TableGrid'
            
    current_directory = os.getcwd()
    file_path = os.path.join(current_directory, "data.docx")
    doc.save(file_path)
    # doc.save("data.docx")
    # os.system('data.docx')
    print(f"Document saved at: {file_path}")
    os.startfile(file_path)

    print("end export")
